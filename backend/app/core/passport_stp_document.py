"""
Формирование паспорта по разделам СТП (для PDF, DOCX, XLSX без шаблона).
"""
from __future__ import annotations

import json
from io import BytesIO
from typing import Any, Dict, List, Optional, Tuple

from openpyxl import Workbook
from openpyxl.styles import Font

from app.core.passport_sections import build_passport_sections, format_formed_at_human

try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
except ImportError:  # pragma: no cover
    Document = None  # type: ignore[misc, assignment]
    qn = None  # type: ignore[misc, assignment]

PASSPORT_FONT_NAME = "Times New Roman"

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    from xml.sax.saxutils import escape
except ImportError:  # pragma: no cover
    SimpleDocTemplate = None  # type: ignore[misc, assignment]


def _sections(envelope: Dict[str, Any], manual: Optional[Dict[str, Any]]) -> List[Dict[str, Any]]:
    return build_passport_sections(envelope, manual)


def _cell_values(row: Dict[str, Any]) -> List[str]:
    return [str(v) if v is not None and v != "" else "—" for v in row.values()]


def _apply_passport_docx_font(document: Any) -> None:
    """Основной шрифт документа — Times New Roman."""
    if qn is None:
        return
    normal = document.styles["Normal"]
    normal.font.name = PASSPORT_FONT_NAME
    normal.font.size = Pt(11)
    rfonts = normal._element.rPr.rFonts  # noqa: SLF001
    rfonts.set(qn("w:eastAsia"), PASSPORT_FONT_NAME)
    rfonts.set(qn("w:ascii"), PASSPORT_FONT_NAME)
    rfonts.set(qn("w:hAnsi"), PASSPORT_FONT_NAME)


def _set_run_times(run: Any, *, bold: bool = False, size_pt: int = 11) -> None:
    run.font.name = PASSPORT_FONT_NAME
    run.font.size = Pt(size_pt)
    run.bold = bold
    if qn is not None:
        r = run._element.get_or_add_rPr()  # noqa: SLF001
        rfonts = r.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), PASSPORT_FONT_NAME)
        rfonts.set(qn("w:ascii"), PASSPORT_FONT_NAME)
        rfonts.set(qn("w:hAnsi"), PASSPORT_FONT_NAME)


def build_docx_from_sections(
    snapshot_envelope: Dict[str, Any],
    title: str,
    manual_sections: Optional[Dict[str, Any]] = None,
) -> bytes:
    if Document is None:
        raise RuntimeError("python-docx не установлен (pip install python-docx)")

    d = Document()
    _apply_passport_docx_font(d)
    h = d.add_heading(title, level=0)
    for run in h.runs:
        _set_run_times(run, bold=True, size_pt=16)

    subtitle = d.add_paragraph()
    r0 = subtitle.add_run("Технический паспорт объекта электросетевого хозяйства\n")
    _set_run_times(r0, bold=True)
    formed = format_formed_at_human(snapshot_envelope.get("formed_at"))
    r1 = subtitle.add_run(f"Дата формирования: {formed}\n")
    _set_run_times(r1)
    r2 = subtitle.add_run(f"Тип объекта: {snapshot_envelope.get('object_type', '—')}\n")
    _set_run_times(r2)
    stp = snapshot_envelope.get("stp_reference") or "—"
    r3 = subtitle.add_run(f"Нормативная база (СТП): {stp}\n")
    _set_run_times(r3)

    for sec in _sections(snapshot_envelope, manual_sections):
        sh = d.add_heading(sec["title"], level=1)
        for run in sh.runs:
            _set_run_times(run, bold=True, size_pt=13)

        if sec.get("rows"):
            t = d.add_table(rows=1, cols=2)
            t.style = "Table Grid"
            t.rows[0].cells[0].text = "Показатель"
            t.rows[0].cells[1].text = "Значение"
            for r in sec["rows"]:
                cells = t.add_row().cells
                cells[0].text = str(r.get("label", ""))
                cells[1].text = str(r.get("value", ""))

        for tbl in sec.get("tables") or []:
            d.add_paragraph()
            p = d.add_paragraph()
            pr = p.add_run(str(tbl.get("title", "")))
            _set_run_times(pr, bold=True)
            columns: List[str] = list(tbl.get("columns") or [])
            rows_data: List[Dict[str, Any]] = list(tbl.get("rows") or [])
            if not columns or not rows_data:
                continue
            table = d.add_table(rows=1, cols=len(columns))
            table.style = "Table Grid"
            for i, col in enumerate(columns):
                table.rows[0].cells[i].text = col
            for row in rows_data:
                vals = _cell_values(row)
                cells = table.add_row().cells
                for i, val in enumerate(vals[: len(columns)]):
                    cells[i].text = val

    bio = BytesIO()
    d.save(bio)
    return bio.getvalue()


def build_pdf_from_sections(
    snapshot_envelope: Dict[str, Any],
    title: str,
    manual_sections: Optional[Dict[str, Any]] = None,
) -> bytes:
    if SimpleDocTemplate is None:
        raise RuntimeError("reportlab не установлен")

    from app.core.passport_export import _p, _p_bold, _p_lines, _register_pdf_cyrillic_font

    font = _register_pdf_cyrillic_font()
    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        rightMargin=1.8 * cm,
        leftMargin=1.8 * cm,
        topMargin=1.8 * cm,
        bottomMargin=1.8 * cm,
        title=(title or "Технический паспорт")[:200],
        author="LEPM",
    )
    story: List[Any] = []
    story.append(_p_bold(title, font, size=14, space_after=8))
    formed = format_formed_at_human(snapshot_envelope.get("formed_at"))
    story.append(
        _p_lines(
            [
                "Технический паспорт объекта электросетевого хозяйства",
                f"Дата формирования: {formed}",
                f"СТП: {snapshot_envelope.get('stp_reference') or '—'}",
            ],
            font,
            size=10,
        )
    )
    story.append(Spacer(1, 0.3 * cm))

    cell_style = ParagraphStyle(
        name="StpCell",
        parent=getSampleStyleSheet()["Normal"],
        fontName=font,
        fontSize=8,
        leading=10,
    )
    head_cell_style = ParagraphStyle(
        name="StpCellHead",
        parent=cell_style,
        fontName=font,
        fontSize=8,
        leading=10,
    )

    def _cell_bold(text: str) -> Paragraph:
        return Paragraph(f"<b>{escape(str(text))}</b>", head_cell_style)

    for sec in _sections(snapshot_envelope, manual_sections):
        story.append(_p_bold(sec["title"], font, size=11, space_after=6))

        if sec.get("rows"):
            kv_data: List[List[Any]] = [
                [
                    _cell_bold("Показатель"),
                    _cell_bold("Значение"),
                ]
            ]
            for r in sec["rows"]:
                kv_data.append(
                    [
                        Paragraph(escape(str(r.get("label", ""))), cell_style),
                        Paragraph(escape(str(r.get("value", ""))), cell_style),
                    ]
                )
            tw = doc.width
            t = Table(kv_data, colWidths=[tw * 0.4, tw * 0.6])
            t.setStyle(
                TableStyle(
                    [
                        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ]
                )
            )
            story.append(t)
            story.append(Spacer(1, 0.25 * cm))

        for tbl in sec.get("tables") or []:
            cols: List[str] = list(tbl.get("columns") or [])
            rows_data = list(tbl.get("rows") or [])
            if not cols or not rows_data:
                continue
            tbl_title = str(tbl.get("title", ""))
            story.append(Paragraph(f"<i>{escape(tbl_title)}</i>", cell_style))
            head = [_cell_bold(c) for c in cols]
            body: List[List[Any]] = [head]
            for row in rows_data[:150]:
                body.append([Paragraph(escape(v), cell_style) for v in _cell_values(row)])
            if len(rows_data) > 150:
                body.append(
                    [
                        Paragraph(escape("…"), cell_style),
                        Paragraph(escape(f"ещё {len(rows_data) - 150} строк"), cell_style),
                    ]
                    + [Paragraph("", cell_style)] * max(0, len(cols) - 2)
                )
            col_w = doc.width / max(len(cols), 1)
            t2 = Table(body, colWidths=[col_w] * len(cols))
            t2.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8e8e8")),
                        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ]
                )
            )
            story.append(t2)
            story.append(Spacer(1, 0.3 * cm))

    doc.build(story)
    return buf.getvalue()


def build_xlsx_from_sections(
    snapshot_envelope: Dict[str, Any],
    title: str,
    manual_sections: Optional[Dict[str, Any]] = None,
) -> bytes:
    """Один лист «Паспорт» — те же разделы, что в PDF и Word."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Паспорт"

    bold = Font(name=PASSPORT_FONT_NAME, bold=True)
    normal = Font(name=PASSPORT_FONT_NAME)
    row_idx = 1

    def next_row() -> int:
        nonlocal row_idx
        r = row_idx
        row_idx += 1
        return r

    def set_cell(r: int, col: int, value: Any, *, font: Optional[Font] = None) -> None:
        c = ws.cell(r, col, value)
        if font is not None:
            c.font = font

    r = next_row()
    set_cell(r, 1, "ТЕХНИЧЕСКИЙ ПАСПОРТ", font=Font(name=PASSPORT_FONT_NAME, bold=True, size=14))
    r = next_row()
    set_cell(r, 1, title, font=Font(name=PASSPORT_FONT_NAME, size=12))
    formed = format_formed_at_human(snapshot_envelope.get("formed_at"))
    for label, val in (
        ("Дата формирования", formed),
        ("Тип объекта", str(snapshot_envelope.get("object_type", ""))),
        ("СТП / норматив", str(snapshot_envelope.get("stp_reference") or "")),
    ):
        r = next_row()
        set_cell(r, 1, label, font=bold)
        set_cell(r, 2, val, font=normal)

    for sec in _sections(snapshot_envelope, manual_sections):
        next_row()
        r = next_row()
        set_cell(r, 1, sec.get("title", ""), font=Font(name=PASSPORT_FONT_NAME, bold=True, size=12))

        if sec.get("rows"):
            r = next_row()
            set_cell(r, 1, "Показатель", font=bold)
            set_cell(r, 2, "Значение", font=bold)
            for item in sec["rows"]:
                r = next_row()
                set_cell(r, 1, item.get("label"), font=normal)
                set_cell(r, 2, item.get("value"), font=normal)

        for tbl in sec.get("tables") or []:
            next_row()
            r = next_row()
            set_cell(r, 1, tbl.get("title", "Таблица"), font=bold)
            cols = list(tbl.get("columns") or [])
            if not cols:
                continue
            r = next_row()
            for ci, col_name in enumerate(cols, start=1):
                set_cell(r, ci, col_name, font=bold)
            for data_row in tbl.get("rows") or []:
                r = next_row()
                for ci, val in enumerate(_cell_values(data_row), start=1):
                    if ci <= len(cols):
                        set_cell(r, ci, val, font=normal)

    ws.column_dimensions["A"].width = 36
    ws.column_dimensions["B"].width = 48
    for col_letter in "CDEFGHIJ":
        ws.column_dimensions[col_letter].width = 18

    bio = BytesIO()
    wb.save(bio)
    return bio.getvalue()
